import tkinter as tk
from tkinter.ttk import Combobox
import docx
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.shared import Pt

gromady= {
	1 : "Антонінська",
	2 : "Берездівська",
	3 : "Білогірська",
	4 : "Війтовецька",
	5 : "Віньковецька",
	6 : "Вовковинецька",
	7 : "Волочиська",
	8 : "Ганнопільська",
	9 : "Гвардійська",
	10 : "Городоцька",
	11 : "Грицівська",
	12 : "Гуківська",
	13 : "Гуменецька",
	14 : "Деражнянська",
	15 : "Дунаєвецька",
	16 : "Жванецька",
	17 : "Закупненська",
	18 : "Заслучненська",
	19 : "Зіньківська",
	20 : "Ізяславська",
	21 : "Кам'янець-Подільська",
	22 : "Китайгородська",
	23 : "Красилівська",
	24 : "Крупецька",
	25 : "Ленковецька",
	26 : "Летичівська",
	27 : "Лісовогринівецька",
	28 : "Маківська",
	29 : "Меджибізька",
	30 : "Миролюбненська",
	31 : "Михайлюцька",
	32 : "Наркевицька",
	33 : "Нетішинська",
	34 : "Новодунаєвецька",
	35 : "Новоушицька",
	36 : "Орининська",
	37 : "Плужненська",
	38 : "Полонська",
	39 : "Понінківська",
	40 : "Розсошанська",
	41 : "Сатанівська",
	42 : "Сахновецька",
	43 : "Славутська",
	44 : "Слобідсько-Кульчієвецька",
	45 : "Смотрицька",
	46 : "Солобковецька",
	47 : "Старокостянтинівська",
	48 : "Староостропільська",
	49 : "Старосинявська",
	50 : "Староушицька",
	51 : "Судилківська",
	52 : "Теофіпольська",
	53 : "Улашанівська",
	54 : "Хмельницька",
	55 : "Чемеровецька",
	56 : "Чорноострівська",
	57 : "Шепетівська",
	58 : "Щиборівська",
	59 : "Ямпільська",
	60 : "Ярмолинецька"
}

def select_gr():
	grom=tup_gr.get()

def select_raj():
	ra=rayon.get()

def get_gr():
	g_a=gromada.get()
	print(gromada)

win=tk.Tk()
win.title('Інвентаризатор')
win.geometry("500x500+250+20")
win.resizable(False,False)

tup_gr = tk.IntVar()
rayon = tk.IntVar()
gromada=tk.IntVar()
g_a=tk.IntVar()
print(gromada)



tk.Label(win,text="Виберіть тип громади:").pack()
tk.Radiobutton(win,text="міська",variable=tup_gr,value=1,command=select_gr()).pack()
tk.Radiobutton(win,text="селищна",variable=tup_gr,value=2,command=select_gr()).pack()
tk.Radiobutton(win,text="сільська",variable=tup_gr,value=3,command=select_gr()).pack()

tk.Label(win,text="Виберіть район:").pack()
tk.Radiobutton(win,text="Шепетівський",variable=rayon,value=1,command=select_raj()).pack()
tk.Radiobutton(win,text="Хмельницький",variable=rayon,value=2,command=select_raj()).pack()
tk.Radiobutton(win,text="Кам'янець-Подільський",variable=rayon,value=3,command=select_raj()).pack()

tk.Label(win,text="Введіть громаду:").pack()
com=Combobox(win,values=list(gromady.values()))
print(com)
com.pack()
#for g_a in sorted(gromady):
#	com=Combobox(win,values=gromady)
#	com.pack()
#	tk.Radiobutton(win,text=gromady[g_a],variable=gromada,value=g_a,command=get_gr()).pack()

tk.Button(win,text="Створити акт інвентаризації", command=get_gr).pack
#print(gr)

win.mainloop()
#gromada=com.value()
#print(gromada)

#gromada = input('Введіть назву громади: ')
name = "C:\Громади\\" + gromada + " ТГ.xlsx"
#print(gromada)
#tup_gr = int(input('Введіть тип громади(1-міська, 2-селищна, 3-сільська): '))
#rayon = int(input('Введіть район(1-Шепетівський, 2-Хмельницький, 3-Кам-Подільський): '))
if tup_gr==1:
	tup="міської"
else:
	if tup_gr==2:
		tup="селищної"
	else:
		if tup_gr==3:
			tup="сільської"	

if rayon==1:
	raj="Шепетівського"
else:
	if rayon==2:
		raj="Хмельницького"
	else:
		if rayon==3:
			raj="Кам'янець-Подільського"	



wb = load_workbook (filename = name)
sheet = wb.active
sheet_ranges = wb['Копія аркуша для акта інвент(2)']

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
win.mainloop()
document = Document("C:\Громади\Акт.docx")

section = document.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(30)
section.right_margin = Mm(10)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

stylet = document.styles['Table Grid']
fontt = stylet.font
fontt.name = 'Times New Roman'
fontt.size = Pt(10)


gr2=gromada[:-1]+"ої"
Dictionary = {"Грицівської": gr2, "селищної": tup, "Шепетівського": raj}

for i in Dictionary:
    for p in document.paragraphs:
        if p.text.find(i)>=0:
            p.text=p.text.replace(i,Dictionary[i])
#for p in document.paragraphs:
#    inline = p.runs
#    for i in range(len(inline)):
#        text = inline[i].text
#        if text in Dictionary.keys():
#            text=text.replace(text,Dictionary[text])
#            inline[i].text = text


#bb = input('Скільки таблиць потрібно створити? ')
bb = int(sheet.cell(row=67, column=1).value)
i =0
col = 2
bb1 = bb + 2
bb2 = bb + 2

#document.paragraphs[1].runs.style.font.size = Pt(12)
#document.paragraphs[1].style.font.size = Pt(12)
#print(document.paragraphs[1].text)


while i < bb1 and col < bb2:
	table = document.add_table(rows=26, cols=2, style = stylet)
	c = table.columns[0].cells
	c[0].text = sheet_ranges['A41'].value
	c[1].text = sheet_ranges['A42'].value
	c[2].text = sheet_ranges['A43'].value
	c[3].text = sheet_ranges['A44'].value
	c[4].text = sheet_ranges['A45'].value
	c[5].text = sheet_ranges['A46'].value
	c[6].text = sheet_ranges['A47'].value
	c[7].text = sheet_ranges['A48'].value
	c[8].text = sheet_ranges['A49'].value
	c[9].text = sheet_ranges['A50'].value
	c[10].text = sheet_ranges['A51'].value
	c[11].text = sheet_ranges['A52'].value
	c[12].text = sheet_ranges['A53'].value
	c[13].text = sheet_ranges['A54'].value
	c[14].text = sheet_ranges['A55'].value
	c[15].text = sheet_ranges['A56'].value
	c[16].text = sheet_ranges['A57'].value
	c[17].text = sheet_ranges['A58'].value
	c[18].text = sheet_ranges['A59'].value
	c[19].text = sheet_ranges['A60'].value
	c[20].text = sheet_ranges['A61'].value
	c[21].text = sheet_ranges['A62'].value
	c[22].text = sheet_ranges['A63'].value
	c[23].text = sheet_ranges['A64'].value
	c[24].text = sheet_ranges['A65'].value
	c[25].text = sheet_ranges['A66'].value
	
	b41=sheet.cell(row=41, column=col).value
	b42=sheet.cell(row=42, column=col).value
	b43=sheet.cell(row=43, column=col).value
	b44=sheet.cell(row=44, column=col).value
	b45=sheet.cell(row=45, column=col).value
	b46=sheet.cell(row=46, column=col).value
	b47=sheet.cell(row=47, column=col).value
	b48=sheet.cell(row=48, column=col).value
	b49=sheet.cell(row=49, column=col).value
	b49=str(b49)
	b50=sheet.cell(row=50, column=col).value
	b50=str(b50)
	b51=sheet.cell(row=51, column=col).value
	b52=sheet.cell(row=52, column=col).value
	b53=sheet.cell(row=53, column=col).value
	b54=sheet.cell(row=54, column=col).value
	b55=sheet.cell(row=55, column=col).value
	b56=sheet.cell(row=56, column=col).value
	b57=sheet.cell(row=57, column=col).value
	b58=sheet.cell(row=58, column=col).value
	b59=sheet.cell(row=59, column=col).value
	b60=sheet.cell(row=60, column=col).value
	b61=sheet.cell(row=61, column=col).value
	b62=sheet.cell(row=62, column=col).value
	b63=sheet.cell(row=63, column=col).value
	b64=sheet.cell(row=64, column=col).value
	b65=sheet.cell(row=65, column=col).value
	b66=sheet.cell(row=66, column=col).value
	
	table.cell (0,1) .text = b41
	table.cell (1,1) .text = b42
	table.cell (2,1) .text = b43
	table.cell (3,1) .text = b44
	table.cell (4,1) .text = b45
	table.cell (5,1) .text = b46
	table.cell (6,1) .text = b47
	table.cell (7,1) .text = b48
	table.cell (8,1) .text = b49
	print(table.cell (8,1).text)
	
	print(table.cell (8,1).text)
	table.cell (9,1) .text = b50
	table.cell (10,1) .text = b51
	table.cell (11,1) .text = b52
	table.cell (12,1) .text = b53
	table.cell (13,1) .text = b54
	table.cell (14,1) .text = b55
	table.cell (15,1) .text = b56
	table.cell (16,1) .text = b57
	table.cell (17,1) .text = b58
	table.cell (18,1) .text = b59
	table.cell (19,1) .text = b60
	table.cell (20,1) .text = b61
	table.cell (21,1) .text = b62
	table.cell (22,1) .text = b63
	table.cell (23,1) .text = b64
	table.cell (24,1) .text = b65
	table.cell (25,1) .text = " "
    
	make_rows_bold(table.rows[0], table.rows[10])

	col=col+1
	i = i + 1
	
	print ( i, col )
	document.add_paragraph(' ')

#print(document.paragraphs[71].text)
#document.paragraphs[71].style.font.size = Pt(10)	


save_name = 'Акт інвентаризації ' + gromada +' ТГ.docx'
print(save_name)
document.save(save_name)

