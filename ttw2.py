from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
gromada = input('Введіть назву громади: ')
name = "С:\Громади\\" + gromada + " ТГ.xlsx"
#name = "..\\" + gromada + " ТГ.xlsx"
print(name)
wb = load_workbook (filename = name)
sheet = wb.active
sheet_ranges = wb['Копія аркуша для акта інвент(2)']

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

document = Document()


bb = input('Скільки таблиць потрібно створити? ')
bb = int(bb)
i =0
col = 2
bb1 = bb + 2
bb2 = bb + 2
while i < bb1 and col < bb2:
	
	table = document.add_table(rows=26, cols=2, style = 'Table Grid')
	c = table.columns[0].cells
	c[0].text = sheet_ranges['A41'].value
	c[1].text = sheet_ranges['A42'].value
	c[2].text = sheet_ranges['A43'].value
	c[3].text = sheet_ranges['A44'].value
	c[4].text = sheet_ranges['A45'].value
	c[5].text = sheet_ranges['A46'].value
	c[6].text = sheet_ranges['A47'].value
	c[7].text = sheet_ranges['A48'].value
	c[8].text = sheet_ranges['A49'].value
	c[9].text = sheet_ranges['A50'].value
	c[10].text = sheet_ranges['A51'].value
	c[11].text = sheet_ranges['A52'].value
	c[12].text = sheet_ranges['A53'].value
	c[13].text = sheet_ranges['A54'].value
	c[14].text = sheet_ranges['A55'].value
	c[15].text = sheet_ranges['A56'].value
	c[16].text = sheet_ranges['A57'].value
	c[17].text = sheet_ranges['A58'].value
	c[18].text = sheet_ranges['A59'].value
	c[19].text = sheet_ranges['A60'].value
	c[20].text = sheet_ranges['A61'].value
	c[21].text = sheet_ranges['A62'].value
	c[22].text = sheet_ranges['A63'].value
	c[23].text = sheet_ranges['A64'].value
	c[24].text = sheet_ranges['A65'].value
	c[25].text = sheet_ranges['A66'].value
	
	b41=sheet.cell(row=41, column=col).value
	b42=sheet.cell(row=42, column=col).value
	b43=sheet.cell(row=43, column=col).value
	b44=sheet.cell(row=44, column=col).value
	b45=sheet.cell(row=45, column=col).value
	b46=sheet.cell(row=46, column=col).value
	b47=sheet.cell(row=47, column=col).value
	b48=sheet.cell(row=48, column=col).value
	b49=sheet.cell(row=49, column=col).value
	b49=str(b49)
	b50=sheet.cell(row=50, column=col).value
	b50=str(b50)
	b51=sheet.cell(row=51, column=col).value
	b52=sheet.cell(row=52, column=col).value
	b53=sheet.cell(row=53, column=col).value
	b54=sheet.cell(row=54, column=col).value
	b55=sheet.cell(row=55, column=col).value
	b56=sheet.cell(row=56, column=col).value
	b57=sheet.cell(row=57, column=col).value
	b58=sheet.cell(row=58, column=col).value
	b59=sheet.cell(row=53, column=col).value
	b60=sheet.cell(row=60, column=col).value
	b61=sheet.cell(row=61, column=col).value
	b62=sheet.cell(row=62, column=col).value
	b63=sheet.cell(row=63, column=col).value
	b64=sheet.cell(row=64, column=col).value
	b65=sheet.cell(row=65, column=col).value
	b66=sheet.cell(row=66, column=col).value
	
	table.cell (0,1) .text = b41
	table.cell (1,1) .text = b42
	table.cell (2,1) .text = b43
	table.cell (3,1) .text = b44
	table.cell (4,1) .text = b45
	table.cell (5,1) .text = b46
	table.cell (6,1) .text = b47
	table.cell (7,1) .text = b48
	table.cell (8,1) .text = b49
	table.cell (9,1) .text = b50
	table.cell (10,1) .text = b51
	table.cell (11,1) .text = b52
	table.cell (12,1) .text = b53
	table.cell (13,1) .text = b54
	table.cell (14,1) .text = b55
	table.cell (15,1) .text = b56
	table.cell (16,1) .text = b57
	table.cell (17,1) .text = b58
	table.cell (18,1) .text = b59
	table.cell (19,1) .text = b60
	table.cell (20,1) .text = b61
	table.cell (21,1) .text = b62
	table.cell (22,1) .text = b63
	table.cell (23,1) .text = b64
	table.cell (24,1) .text = b65
	table.cell (25,1) .text = " "

	make_rows_bold(table.rows[0], table.rows[10])

	col=col+1
	i = i + 1
	
	print ( i, col )
	document.add_paragraph(' ')
	
	



save_name = 'Акт ' + gromada +' ТГ.docx'
document.save(save_name)