from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches

wb = load_workbook(filename = 'C:\Table_to_word\Table_to_word\Михайлюцька ТГ.xlsx')
sheet_ranges = wb['Копія аркуша для акта інвент(2)']
print(sheet_ranges['A41'].value)

document = Document()


#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True
a="A"+41
print(a)
a41 = sheet_ranges['A41'].value

table = document.add_table(rows=26, cols=2)
document.add_page_break()

table.cell (0,0) .text = a41
table.cell (1,0) .text = sheet_ranges['A42'].value
table.cell (2,0) .text = sheet_ranges['A43'].value
table.cell (3,0) .text = sheet_ranges['A44'].value
table.cell (4,0) .text = sheet_ranges['A45'].value
table.cell (5,0) .text = sheet_ranges['A46'].value
table.cell (6,0) .text = sheet_ranges['A47'].value
table.cell (7,0) .text = sheet_ranges['A48'].value
table.cell (8,0) .text = sheet_ranges['A49'].value
table.cell (9,0) .text = sheet_ranges['A50'].value
table.cell (10,0) .text = sheet_ranges['A51'].value
table.cell (11,0) .text = sheet_ranges['A52'].value
table.cell (12,0) .text = sheet_ranges['A53'].value
table.cell (13,0) .text = sheet_ranges['A54'].value
table.cell (14,0) .text = sheet_ranges['A55'].value
table.cell (15,0) .text = sheet_ranges['A56'].value
table.cell (16,0) .text = sheet_ranges['A57'].value
table.cell (17,0) .text = sheet_ranges['A58'].value
table.cell (18,0) .text = sheet_ranges['A59'].value
table.cell (19,0) .text = sheet_ranges['A60'].value
table.cell (20,0) .text = sheet_ranges['A61'].value
table.cell (21,0) .text = sheet_ranges['A62'].value
table.cell (22,0) .text = sheet_ranges['A63'].value
table.cell (23,0) .text = sheet_ranges['A64'].value
table.cell (24,0) .text = sheet_ranges['A65'].value
table.cell (25,0) .text = sheet_ranges['A66'].value

table.cell (0,1) .text = sheet_ranges['B41'].value
table.cell (1,1) .text = sheet_ranges['B42'].value
table.cell (2,1) .text = sheet_ranges['B43'].value
table.cell (3,1) .text = sheet_ranges['B44'].value
table.cell (4,1) .text = sheet_ranges['B45'].value
table.cell (5,1) .text = sheet_ranges['B46'].value
        
document.add_page_break()

document.add_paragraph(sheet_ranges['A41'].value)
document.add_page_break()

document.save('demo.docx')